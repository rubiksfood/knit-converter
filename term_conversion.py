from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from typing import Dict, List


from heading_safe import replace_terms_heading_safe
from text_processing import replace_terms_in_text


def _run_has_picture(run: Run) -> bool:
    # Return true if the run contains an inline picture/drawing

    r = run._r  # lxml element
    # Covers modern drawings and older images
    return bool(r.xpath(".//w:drawing | .//w:pict"))


def _editable_runs(paragraph: Paragraph) -> List[Run]:
    # Edit runs that do not contain images.

    return [run for run in paragraph.runs if not _run_has_picture(run)]


def _replace_in_paragraph_preserve_runs(paragraph: Paragraph, terms: Dict[str, str]) -> bool:
    # Replace terms in a paragraph while preserving run formatting and images

    runs = _editable_runs(paragraph)
    if not runs:
        return False

    # 1) Concatenate all editable run text
    original_parts = [run.text for run in runs]
    original_text = "".join(original_parts)

    # 2) Run replacement on that string
    new_text = replace_terms_in_text(original_text, terms)
    if new_text == original_text:
        return False

    # 3) Write the updated text back into the same runs (same formatting), 
    # with slices of new_text matching original run lengths.
    lengths = [len(t) for t in original_parts]
    idx = 0
    for run, n in zip(runs, lengths):
        run.text = new_text[idx : idx + n]
        idx += n

    # If new_text is longer than original, append leftover to the last editable run.
    # If shorter, the remaining runs become empty.
    if idx < len(new_text):
        runs[-1].text += new_text[idx:]

    # If changes are made, return True
    return True


def convert_terms_in_docx(doc_path: str, terms: Dict[str, str]) -> None:
    # Convert terms in a DOCX while preserving formatting and images as much as possible.

    doc = Document(doc_path)
    changed = 0

    # Normal paragraphs
    for para in doc.paragraphs:
        has_changed = replace_terms_heading_safe(para, terms, _replace_in_paragraph_preserve_runs)
        if has_changed:
            changed += 1

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    has_changed = replace_terms_heading_safe(
                        para, terms, _replace_in_paragraph_preserve_runs
                    )
                    if has_changed:
                        changed += 1

    doc.save(doc_path)